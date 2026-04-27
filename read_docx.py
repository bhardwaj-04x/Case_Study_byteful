import docx

doc = docx.Document(r'c:\Users\prana\OneDrive\Desktop\CaseStudyByteful\case-study.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')

for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        print(f'{i+1}: {para.text[:150]}...')